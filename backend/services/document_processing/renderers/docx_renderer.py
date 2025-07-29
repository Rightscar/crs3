"""
DOCX Document Renderer
=====================

Renderer for Microsoft Word documents using python-docx.
"""

import io
import logging
from typing import Optional, Union, BinaryIO, List, Dict, Any
from datetime import datetime

try:
    from docx import Document as DocxDocument
    from docx.opc.exceptions import PackageNotFoundError
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    logging.warning("python-docx not available - DOCX support disabled")

from backend.services.document_processing.universal_reader import (
    BaseRenderer, DocumentMetadata, DocumentPage, DocumentFormat, TextBlock
)
from backend.core.logging import get_logger

logger = get_logger(__name__)


class DocxRenderer(BaseRenderer):
    """DOCX document renderer using python-docx"""
    
    async def load(self, file_data: Union[bytes, str, BinaryIO]) -> bool:
        """Load DOCX document"""
        if not DOCX_AVAILABLE:
            logger.error("python-docx not available for DOCX rendering")
            return False
        
        try:
            if isinstance(file_data, str):
                # File path
                self.document = DocxDocument(file_data)
            elif isinstance(file_data, bytes):
                # Bytes data
                self.document = DocxDocument(io.BytesIO(file_data))
            else:
                # File-like object
                self.document = DocxDocument(file_data)
            
            await self._extract_metadata()
            return True
        except PackageNotFoundError:
            logger.error("Invalid DOCX file format")
            return False
        except Exception as e:
            logger.error(f"Failed to load DOCX: {e}")
            return False
    
    async def _extract_metadata(self):
        """Extract DOCX metadata"""
        if not self.document:
            return
        
        core_props = self.document.core_properties
        
        self.metadata = DocumentMetadata(
            title=core_props.title or "Unknown Document",
            author=core_props.author or "Unknown Author",
            subject=core_props.subject or "",
            creator=core_props.author or "",
            creation_date=core_props.created,
            modification_date=core_props.modified,
            page_count=len(self.document.paragraphs),  # Approximate
            format=DocumentFormat.DOCX,
            keywords=core_props.keywords.split(',') if core_props.keywords else [],
            language=core_props.language or "en"
        )
    
    def get_page_count(self) -> int:
        """Get approximate page count (based on paragraphs)"""
        if not self.document:
            return 0
        # Rough estimate: 3 paragraphs per page
        return max(1, len(self.document.paragraphs) // 3)
    
    async def render_page(self, page_number: int, scale: float = 1.0) -> Optional[DocumentPage]:
        """
        Render a specific page (approximated by paragraph ranges).
        Note: DOCX doesn't have true pages, so we simulate them.
        """
        if not self.document:
            return None
        
        # Calculate paragraph range for this "page"
        paragraphs_per_page = 3
        start_idx = (page_number - 1) * paragraphs_per_page
        end_idx = start_idx + paragraphs_per_page
        
        paragraphs = list(self.document.paragraphs)
        if start_idx >= len(paragraphs):
            return None
        
        page_paragraphs = paragraphs[start_idx:end_idx]
        
        # Extract text content
        text_content = ""
        text_blocks = []
        y_position = 0
        
        for para in page_paragraphs:
            para_text = para.text.strip()
            if para_text:
                # Create text block
                text_blocks.append(TextBlock(
                    text=para_text,
                    x=0,
                    y=y_position,
                    width=600,  # Standard page width
                    height=20 * (1 + para_text.count('\n')),  # Estimate height
                    font_size=self._get_font_size(para),
                    font_name=self._get_font_name(para)
                ))
                text_content += para_text + "\n\n"
                y_position += 30  # Spacing between paragraphs
        
        # Extract tables on this page
        tables_text = await self._extract_tables_in_range(start_idx, end_idx)
        if tables_text:
            text_content += "\n" + tables_text
        
        return DocumentPage(
            page_number=page_number,
            text_content=text_content.strip(),
            width=600,
            height=800,
            text_blocks=text_blocks
        )
    
    def _get_font_size(self, paragraph) -> Optional[float]:
        """Get font size from paragraph"""
        try:
            if paragraph.runs and paragraph.runs[0].font.size:
                return paragraph.runs[0].font.size.pt
        except:
            pass
        return 11.0  # Default
    
    def _get_font_name(self, paragraph) -> Optional[str]:
        """Get font name from paragraph"""
        try:
            if paragraph.runs and paragraph.runs[0].font.name:
                return paragraph.runs[0].font.name
        except:
            pass
        return "Calibri"  # Default
    
    async def _extract_tables_in_range(self, start_idx: int, end_idx: int) -> str:
        """Extract tables that appear in the paragraph range"""
        tables_text = ""
        
        # Simple approach: extract all tables
        # In reality, we'd need to determine table positions relative to paragraphs
        for table in self.document.tables:
            table_text = "Table:\n"
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells)
                table_text += row_text + "\n"
            tables_text += table_text + "\n"
        
        return tables_text.strip()
    
    async def extract_text(self, page_number: Optional[int] = None) -> str:
        """Extract text from DOCX"""
        if not self.document:
            return ""
        
        try:
            if page_number:
                # Extract from specific "page"
                page = await self.render_page(page_number)
                return page.text_content if page else ""
            else:
                # Extract all text
                text_parts = []
                
                # Extract paragraphs
                for para in self.document.paragraphs:
                    if para.text.strip():
                        text_parts.append(para.text.strip())
                
                # Extract tables
                for table in self.document.tables:
                    table_text = []
                    for row in table.rows:
                        row_text = " | ".join(cell.text.strip() for cell in row.cells)
                        table_text.append(row_text)
                    if table_text:
                        text_parts.append("\n".join(table_text))
                
                return "\n\n".join(text_parts)
        except Exception as e:
            logger.error(f"Failed to extract text: {e}")
            return ""
    
    async def get_metadata(self) -> DocumentMetadata:
        """Get DOCX metadata"""
        return self.metadata or DocumentMetadata(format=DocumentFormat.DOCX)
    
    async def search(self, query: str, case_sensitive: bool = False) -> List[Dict[str, Any]]:
        """Search for text in DOCX"""
        if not self.document:
            return []
        
        results = []
        query_lower = query if case_sensitive else query.lower()
        
        try:
            # Search in paragraphs
            for i, para in enumerate(self.document.paragraphs):
                text = para.text
                search_text = text if case_sensitive else text.lower()
                
                if query_lower in search_text:
                    # Find all occurrences
                    start = 0
                    while True:
                        pos = search_text.find(query_lower, start)
                        if pos == -1:
                            break
                        
                        results.append({
                            "page": i // 3 + 1,  # Approximate page
                            "paragraph": i,
                            "text": query,
                            "context": text[max(0, pos-50):pos+50+len(query)],
                            "position": pos
                        })
                        start = pos + 1
            
            # Search in tables
            for table_idx, table in enumerate(self.document.tables):
                for row_idx, row in enumerate(table.rows):
                    for cell_idx, cell in enumerate(row.cells):
                        text = cell.text
                        search_text = text if case_sensitive else text.lower()
                        
                        if query_lower in search_text:
                            results.append({
                                "page": -1,  # Tables don't have page numbers
                                "table": table_idx,
                                "row": row_idx,
                                "cell": cell_idx,
                                "text": query,
                                "context": text
                            })
        except Exception as e:
            logger.error(f"Search failed: {e}")
        
        return results
    
    async def close(self):
        """Clean up resources"""
        self.document = None