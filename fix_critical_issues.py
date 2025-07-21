#!/usr/bin/env python3
"""
Critical Issues Fix Script
=========================

This script fixes the most critical issues found in the codebase
that would prevent deployment or cause runtime failures.
"""

import os
import sys
import logging
from pathlib import Path

# Configure logging
logging.basicConfig(level=logging.INFO, format='%(levelname)s: %(message)s')
logger = logging.getLogger(__name__)

def create_missing_modules():
    """Create minimal implementations of missing modules"""
    
    modules_dir = Path("modules")
    
    # Create basic DOCX renderer
    docx_renderer_content = '''"""
Basic DOCX Renderer
==================
Minimal implementation for DOCX file support.
"""

import logging
from typing import Optional, Union
from .universal_document_reader import BaseRenderer, DocumentPage, DocumentMetadata

logger = logging.getLogger(__name__)

class DocxRenderer(BaseRenderer):
    """Basic DOCX document renderer"""
    
    def __init__(self):
        super().__init__()
        self.document = None
    
    def load(self, file_data: Union[bytes, str]) -> bool:
        """Load DOCX document"""
        try:
            from docx import Document
            if isinstance(file_data, bytes):
                import io
                self.document = Document(io.BytesIO(file_data))
            else:
                self.document = Document(file_data)
            return True
        except Exception as e:
            logger.error(f"Failed to load DOCX: {e}")
            return False
    
    def get_page_count(self) -> int:
        """Get page count (DOCX doesn't have explicit pages)"""
        return 1 if self.document else 0
    
    def get_metadata(self) -> DocumentMetadata:
        """Extract DOCX metadata"""
        if not self.document:
            return DocumentMetadata()
        
        core_props = self.document.core_properties
        return DocumentMetadata(
            title=core_props.title or "Word Document",
            author=core_props.author or "",
            subject=core_props.subject or "",
            format='docx',
            page_count=1
        )
    
    def render_page(self, page_number: int, zoom: float = 1.0) -> Optional[DocumentPage]:
        """Render DOCX as single page"""
        if not self.document or page_number != 1:
            return None
        
        # Extract all text
        text_content = "\\n".join([paragraph.text for paragraph in self.document.paragraphs])
        
        return DocumentPage(
            page_number=1,
            text_content=text_content,
            width=800,
            height=1000
        )
    
    def extract_text(self, page_number: int) -> str:
        """Extract text from DOCX"""
        if not self.document:
            return ""
        
        return "\\n".join([paragraph.text for paragraph in self.document.paragraphs])
    
    def search_text(self, query: str, page_number: int = None) -> list:
        """Search text in DOCX"""
        results = []
        text = self.extract_text(1)
        
        if query.lower() in text.lower():
            results.append({
                'page': 1,
                'text': query,
                'context': text[:200] + "..."
            })
        
        return results
'''
    
    # Create basic EPUB renderer
    epub_renderer_content = '''"""
Basic EPUB Renderer
==================
Minimal implementation for EPUB file support.
"""

import logging
from typing import Optional, Union, List, Dict
from .universal_document_reader import BaseRenderer, DocumentPage, DocumentMetadata

logger = logging.getLogger(__name__)

class EpubRenderer(BaseRenderer):
    """Basic EPUB document renderer"""
    
    def __init__(self):
        super().__init__()
        self.book = None
        self.chapters = []
    
    def load(self, file_data: Union[bytes, str]) -> bool:
        """Load EPUB document"""
        try:
            import ebooklib
            from ebooklib import epub
            
            if isinstance(file_data, bytes):
                import io
                self.book = epub.read_epub(io.BytesIO(file_data))
            else:
                self.book = epub.read_epub(file_data)
            
            # Extract chapters
            self.chapters = []
            for item in self.book.get_items():
                if item.get_type() == ebooklib.ITEM_DOCUMENT:
                    self.chapters.append(item)
            
            return True
        except Exception as e:
            logger.error(f"Failed to load EPUB: {e}")
            return False
    
    def get_page_count(self) -> int:
        """Get chapter count as page count"""
        return len(self.chapters)
    
    def get_metadata(self) -> DocumentMetadata:
        """Extract EPUB metadata"""
        if not self.book:
            return DocumentMetadata()
        
        title = ""
        author = ""
        
        for meta in self.book.get_metadata('DC', 'title'):
            title = meta[0]
            break
            
        for meta in self.book.get_metadata('DC', 'creator'):
            author = meta[0]
            break
        
        return DocumentMetadata(
            title=title or "EPUB Document",
            author=author or "",
            format='epub',
            page_count=len(self.chapters)
        )
    
    def render_page(self, page_number: int, zoom: float = 1.0) -> Optional[DocumentPage]:
        """Render EPUB chapter as page"""
        if not self.book or page_number < 1 or page_number > len(self.chapters):
            return None
        
        chapter = self.chapters[page_number - 1]
        
        # Extract text from HTML content
        import re
        from html import unescape
        
        content = chapter.get_content().decode('utf-8')
        # Strip HTML tags
        text_content = re.sub(r'<[^>]+>', '', content)
        text_content = unescape(text_content)
        
        return DocumentPage(
            page_number=page_number,
            text_content=text_content,
            width=600,
            height=800
        )
    
    def extract_text(self, page_number: int) -> str:
        """Extract text from EPUB chapter"""
        page = self.render_page(page_number)
        return page.text_content if page else ""
    
    def search_text(self, query: str, page_number: int = None) -> list:
        """Search text in EPUB"""
        results = []
        
        if page_number:
            pages_to_search = [page_number]
        else:
            pages_to_search = range(1, len(self.chapters) + 1)
        
        for page_num in pages_to_search:
            text = self.extract_text(page_num)
            if query.lower() in text.lower():
                results.append({
                    'page': page_num,
                    'text': query,
                    'context': text[:200] + "..."
                })
        
        return results
'''
    
    # Write the files
    docx_file = modules_dir / "docx_renderer.py"
    epub_file = modules_dir / "epub_renderer.py"
    
    try:
        with open(docx_file, 'w') as f:
            f.write(docx_renderer_content)
        logger.info(f"Created {docx_file}")
        
        with open(epub_file, 'w') as f:
            f.write(epub_renderer_content)
        logger.info(f"Created {epub_file}")
        
        return True
    except Exception as e:
        logger.error(f"Failed to create missing modules: {e}")
        return False

def create_health_check():
    """Create a health check endpoint for Streamlit"""
    
    health_check_content = '''"""
Health Check Module
==================
Provides health check functionality for deployment platforms.
"""

import streamlit as st
import os
import sys
from datetime import datetime

def check_health():
    """Basic health check function"""
    try:
        # Check if main modules can be imported
        from modules.universal_document_reader import UniversalDocumentReader
        from modules.intelligent_processor import IntelligentProcessor
        
        # Check environment variables
        required_env_vars = ['STREAMLIT_SERVER_PORT']
        missing_vars = [var for var in required_env_vars if not os.getenv(var)]
        
        health_status = {
            'status': 'healthy',
            'timestamp': datetime.now().isoformat(),
            'python_version': sys.version,
            'modules_loaded': True,
            'missing_env_vars': missing_vars
        }
        
        if missing_vars:
            health_status['status'] = 'degraded'
        
        return health_status
        
    except Exception as e:
        return {
            'status': 'unhealthy',
            'timestamp': datetime.now().isoformat(),
            'error': str(e)
        }

# Add health check route (if running in debug mode)
if __name__ == "__main__":
    import json
    print(json.dumps(check_health(), indent=2))
'''
    
    try:
        with open("health_check.py", 'w') as f:
            f.write(health_check_content)
        logger.info("Created health_check.py")
        return True
    except Exception as e:
        logger.error(f"Failed to create health check: {e}")
        return False

def create_streamlit_config():
    """Create Streamlit configuration for production"""
    
    config_dir = Path(".streamlit")
    config_dir.mkdir(exist_ok=True)
    
    config_content = '''[server]
port = 8501
address = "0.0.0.0"
maxUploadSize = 50
enableCORS = false
enableXsrfProtection = true

[browser]
gatherUsageStats = false

[theme]
primaryColor = "#667eea"
backgroundColor = "#0f0f23"
secondaryBackgroundColor = "#1a1a2e"
textColor = "#ffffff"

[logger]
level = "info"
'''
    
    try:
        config_file = config_dir / "config.toml"
        with open(config_file, 'w') as f:
            f.write(config_content)
        logger.info(f"Created {config_file}")
        return True
    except Exception as e:
        logger.error(f"Failed to create Streamlit config: {e}")
        return False

def fix_requirements():
    """Add version pins to requirements.txt for stability"""
    
    try:
        # Read current requirements
        with open("requirements.txt", 'r') as f:
            lines = f.readlines()
        
        # Add specific version pins for critical packages
        version_pins = {
            'streamlit': '>=1.29.0,<2.0.0',
            'pandas': '>=2.1.0,<3.0.0',
            'numpy': '>=1.25.0,<2.0.0',
            'openai': '>=1.0.0,<2.0.0',
            'PyMuPDF': '>=1.23.0,<2.0.0'
        }
        
        updated_lines = []
        for line in lines:
            package_name = line.split('>=')[0].split('==')[0].strip()
            if package_name in version_pins:
                updated_lines.append(f"{package_name}{version_pins[package_name]}\\n")
            else:
                updated_lines.append(line)
        
        # Write updated requirements
        with open("requirements.txt", 'w') as f:
            f.writelines(updated_lines)
        
        logger.info("Updated requirements.txt with version pins")
        return True
    except Exception as e:
        logger.error(f"Failed to update requirements: {e}")
        return False

def main():
    """Main function to run all fixes"""
    logger.info("Starting critical issues fix...")
    
    fixes_applied = 0
    total_fixes = 4
    
    # Create missing modules
    if create_missing_modules():
        fixes_applied += 1
        logger.info("✅ Created missing renderer modules")
    else:
        logger.error("❌ Failed to create missing modules")
    
    # Create health check
    if create_health_check():
        fixes_applied += 1
        logger.info("✅ Created health check module")
    else:
        logger.error("❌ Failed to create health check")
    
    # Create Streamlit config
    if create_streamlit_config():
        fixes_applied += 1
        logger.info("✅ Created Streamlit configuration")
    else:
        logger.error("❌ Failed to create Streamlit config")
    
    # Fix requirements
    if fix_requirements():
        fixes_applied += 1
        logger.info("✅ Updated requirements.txt")
    else:
        logger.error("❌ Failed to update requirements")
    
    logger.info(f"Applied {fixes_applied}/{total_fixes} fixes")
    
    if fixes_applied == total_fixes:
        logger.info("🎉 All critical issues fixed successfully!")
        print("\\n" + "="*50)
        print("NEXT STEPS:")
        print("1. Set your OpenAI API key in .env file")
        print("2. Install dependencies: pip install -r requirements.txt")
        print("3. Download spaCy model: python -m spacy download en_core_web_sm")
        print("4. Test the application: streamlit run app.py")
        print("="*50)
    else:
        logger.warning("Some fixes failed. Check the logs above.")
        sys.exit(1)

if __name__ == "__main__":
    main()