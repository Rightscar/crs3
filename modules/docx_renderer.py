"""
DOCX Renderer Module
===================

Handles rendering and processing of DOCX (Microsoft Word) documents.
This module was missing and causing import failures.

Features:
- DOCX document loading and rendering
- Text extraction from DOCX files
- Page-by-page content extraction
- Error handling for corrupted DOCX files
"""

import logging
from typing import Dict, List, Any, Optional, Union, Tuple
import io

# Optional imports with fallbacks
try:
    from docx import Document as DocxDocument
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    logging.warning("python-docx not available - DOCX rendering will be limited")

logger = logging.getLogger(__name__)

class DocxRenderer:
    """DOCX document renderer with error handling"""
    
    def __init__(self):
        self.document = None
        self.docx_available = DOCX_AVAILABLE
        
    def load(self, file_data: Union[bytes, io.IOBase]) -> Dict[str, Any]:
        """
        Load DOCX document from file data
        
        Args:
            file_data: DOCX file content as bytes or file-like object
            
        Returns:
            Dict with success status and document info
        """
        try:
            if not self.docx_available:
                return {
                    'success': False,
                    'error': 'python-docx not available - cannot process DOCX files'
                }
            
            # Load document
            if isinstance(file_data, bytes):
                file_stream = io.BytesIO(file_data)
            else:
                file_stream = file_data
                
            self.document = DocxDocument(file_stream)
            
            # Extract basic metadata
            metadata = self._extract_metadata()
            
            return {
                'success': True,
                'total_pages': len(self.document.paragraphs),  # Approximate page count
                'metadata': metadata
            }
            
        except Exception as e:
            logger.error(f"DOCX loading error: {e}")
            return {
                'success': False,
                'error': f'Failed to load DOCX: {str(e)}'
            }
    
    def get_page_count(self) -> int:
        """Get total number of pages (approximated for DOCX)"""
        if not self.document:
            return 0
        return max(1, len(self.document.paragraphs) // 10)  # Rough estimate
    
    def extract_page_text(self, page_number: int) -> str:
        """
        Extract text from a specific page
        
        Args:
            page_number: Page number (1-indexed)
            
        Returns:
            Text content of the page
        """
        if not self.document:
            return ""
            
        try:
            # For DOCX, we approximate pages by paragraph groups
            paragraphs_per_page = 10
            start_idx = (page_number - 1) * paragraphs_per_page
            end_idx = start_idx + paragraphs_per_page
            
            paragraphs = self.document.paragraphs[start_idx:end_idx]
            # Safe text extraction from paragraphs
            text_parts = []
            for p in paragraphs:
                if p and hasattr(p, 'text'):
                    text_parts.append(p.text if p.text else "")
            text_content = '\n'.join(text_parts)
            
            return text_content
            
        except Exception as e:
            logger.error(f"Text extraction error for page {page_number}: {e}")
            return ""
    
    def extract_all_text(self) -> str:
        """Extract all text from the document"""
        if not self.document:
            return ""
            
        try:
            all_text = []
            for paragraph in self.document.paragraphs:
                all_text.append(paragraph.text)
            
            return '\n'.join(all_text)
            
        except Exception as e:
            logger.error(f"Full text extraction error: {e}")
            return ""
    
    def _extract_metadata(self) -> Dict[str, Any]:
        """Extract metadata from DOCX document"""
        if not self.document:
            return {}
            
        try:
            core_props = self.document.core_properties
            
            metadata = {
                'title': getattr(core_props, 'title', None) or 'Unknown',
                'author': getattr(core_props, 'author', None) or 'Unknown',
                'subject': getattr(core_props, 'subject', None) or '',
                'created': getattr(core_props, 'created', None),
                'modified': getattr(core_props, 'modified', None),
                'paragraph_count': len(self.document.paragraphs)
            }
            
            return metadata
            
        except Exception as e:
            logger.error(f"Metadata extraction error: {e}")
            return {'error': str(e)}
    
    def get_table_of_contents(self) -> List[Dict[str, Any]]:
        """Extract table of contents (basic implementation for DOCX)"""
        if not self.document:
            return []
            
        try:
            toc = []
            for i, paragraph in enumerate(self.document.paragraphs):
                # Look for heading-style paragraphs
                if paragraph.style and 'heading' in paragraph.style.name.lower():
                    toc.append({
                        'title': paragraph.text,
                        'page': (i // 10) + 1,  # Approximate page
                        'level': int(paragraph.style.name[-1]) if paragraph.style.name[-1].isdigit() else 1
                    })
            
            return toc[:20]  # Limit to first 20 headings
            
        except Exception as e:
            logger.error(f"TOC extraction error: {e}")
            return []
    
    def render_page_image(self, page_number: int) -> Optional[bytes]:
        """
        Render page as image (not supported for DOCX)
        
        Returns:
            None (DOCX doesn't support direct image rendering)
        """
        logger.warning("Image rendering not supported for DOCX files")
        return None
    
    def search_text(self, query: str) -> List[Dict[str, Any]]:
        """
        Search for text in the document
        
        Args:
            query: Search query
            
        Returns:
            List of search results with context
        """
        if not self.document or not query:
            return []
            
        try:
            results = []
            query_lower = query.lower()
            
            for i, paragraph in enumerate(self.document.paragraphs):
                text = paragraph.text
                if query_lower in text.lower():
                    results.append({
                        'page': (i // 10) + 1,  # Approximate page
                        'text': text,
                        'context': text[:200],  # First 200 chars as context
                        'match_type': 'exact'
                    })
            
            return results[:50]  # Limit results
            
        except Exception as e:
            logger.error(f"Search error: {e}")
            return []
