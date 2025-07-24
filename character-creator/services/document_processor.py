"""
Document Processing Service
===========================

Process various document formats for character extraction.
"""

import os
import hashlib
from pathlib import Path
from typing import Optional, Dict, Any, List
import PyPDF2
import docx
import ebooklib
from ebooklib import epub
import fitz  # PyMuPDF
from PIL import Image
import pytesseract

from config.settings import settings, UPLOAD_DIR
from config.logging_config import logger
from core.exceptions import DocumentProcessingError
from core.security import security
from core.models import DocumentReference

class DocumentProcessor:
    """Process documents for character extraction"""
    
    def __init__(self):
        self.supported_formats = {
            '.pdf': self._process_pdf,
            '.txt': self._process_text,
            '.md': self._process_text,
            '.docx': self._process_docx,
            '.doc': self._process_docx,
            '.epub': self._process_epub,
            '.rtf': self._process_rtf,
            '.html': self._process_html
        }
        
    def process_document(self, file_path: str, filename: str) -> Dict[str, Any]:
        """
        Process uploaded document
        
        Args:
            file_path: Path to uploaded file
            filename: Original filename
            
        Returns:
            Document data including text and metadata
        """
        try:
            # Validate file
            security.validate_filename(filename)
            
            # Get file extension
            ext = Path(filename).suffix.lower()
            
            if ext not in self.supported_formats:
                raise DocumentProcessingError(
                    f"Unsupported file format: {ext}",
                    details={'supported': list(self.supported_formats.keys())}
                )
            
            # Process based on format
            processor = self.supported_formats[ext]
            text, metadata = processor(file_path)
            
            # Calculate document hash
            with open(file_path, 'rb') as f:
                content_hash = security.hash_content(f.read())
            
            # Count words and estimate pages
            word_count = len(text.split())
            page_count = metadata.get('page_count', word_count // 250)  # Estimate if not available
            
            # Create document reference
            doc_ref = DocumentReference(
                filename=filename,
                filepath=file_path,
                content_hash=content_hash,
                total_pages=page_count,
                word_count=word_count,
                metadata=metadata
            )
            
            logger.info(f"Processed document: {filename} ({word_count} words, {page_count} pages)")
            
            return {
                'text': text,
                'document_reference': doc_ref,
                'metadata': metadata,
                'success': True
            }
            
        except Exception as e:
            logger.error(f"Error processing document {filename}: {e}")
            raise DocumentProcessingError(
                f"Failed to process document: {str(e)}",
                details={'filename': filename}
            )
    
    def _process_pdf(self, file_path: str) -> tuple[str, Dict]:
        """Process PDF files"""
        text = ""
        metadata = {'page_count': 0, 'has_images': False}
        
        try:
            # Try PyMuPDF first (better for complex PDFs)
            doc = fitz.open(file_path)
            
            for page_num, page in enumerate(doc):
                # Extract text
                page_text = page.get_text()
                
                # If no text, try OCR on images
                if not page_text.strip():
                    pix = page.get_pixmap()
                    img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
                    
                    try:
                        page_text = pytesseract.image_to_string(img)
                        metadata['has_images'] = True
                    except:
                        logger.warning(f"OCR failed for page {page_num + 1}")
                
                text += page_text + "\n"
            
            metadata['page_count'] = len(doc)
            doc.close()
            
        except Exception as e:
            logger.warning(f"PyMuPDF failed, trying PyPDF2: {e}")
            
            # Fallback to PyPDF2
            with open(file_path, 'rb') as f:
                reader = PyPDF2.PdfReader(f)
                metadata['page_count'] = len(reader.pages)
                
                for page in reader.pages:
                    text += page.extract_text() + "\n"
        
        return text, metadata
    
    def _process_text(self, file_path: str) -> tuple[str, Dict]:
        """Process plain text files"""
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
            text = f.read()
        
        metadata = {
            'page_count': len(text) // 3000 + 1,  # Rough estimate
            'encoding': 'utf-8'
        }
        
        return text, metadata
    
    def _process_docx(self, file_path: str) -> tuple[str, Dict]:
        """Process Word documents"""
        doc = docx.Document(file_path)
        
        text = ""
        metadata = {
            'paragraph_count': len(doc.paragraphs),
            'table_count': len(doc.tables)
        }
        
        # Extract text from paragraphs
        for para in doc.paragraphs:
            text += para.text + "\n"
        
        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text += cell.text + " "
            text += "\n"
        
        # Estimate page count
        metadata['page_count'] = len(text) // 3000 + 1
        
        return text, metadata
    
    def _process_epub(self, file_path: str) -> tuple[str, Dict]:
        """Process EPUB files"""
        book = epub.read_epub(file_path)
        
        text = ""
        metadata = {
            'title': book.get_metadata('DC', 'title'),
            'author': book.get_metadata('DC', 'creator'),
            'chapter_count': 0
        }
        
        # Extract text from all items
        for item in book.get_items():
            if item.get_type() == ebooklib.ITEM_DOCUMENT:
                # Decode and clean HTML
                content = item.get_content().decode('utf-8', errors='ignore')
                
                # Simple HTML stripping
                import re
                clean_text = re.sub('<[^<]+?>', '', content)
                text += clean_text + "\n"
                metadata['chapter_count'] += 1
        
        metadata['page_count'] = len(text) // 3000 + 1
        
        return text, metadata
    
    def _process_rtf(self, file_path: str) -> tuple[str, Dict]:
        """Process RTF files"""
        try:
            from striprtf.striprtf import rtf_to_text
            
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                rtf_content = f.read()
            
            text = rtf_to_text(rtf_content)
            
        except ImportError:
            # Fallback: treat as text
            logger.warning("striprtf not installed, treating RTF as plain text")
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                text = f.read()
        
        metadata = {'page_count': len(text) // 3000 + 1}
        
        return text, metadata
    
    def _process_html(self, file_path: str) -> tuple[str, Dict]:
        """Process HTML files"""
        from bs4 import BeautifulSoup
        
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
            html_content = f.read()
        
        soup = BeautifulSoup(html_content, 'html.parser')
        
        # Remove script and style elements
        for script in soup(["script", "style"]):
            script.decompose()
        
        # Get text
        text = soup.get_text()
        
        # Clean up whitespace
        lines = (line.strip() for line in text.splitlines())
        chunks = (phrase.strip() for line in lines for phrase in line.split("  "))
        text = '\n'.join(chunk for chunk in chunks if chunk)
        
        metadata = {
            'page_count': len(text) // 3000 + 1,
            'has_images': len(soup.find_all('img')) > 0
        }
        
        return text, metadata
    
    def extract_chapters(self, text: str) -> List[Dict[str, Any]]:
        """
        Extract chapters from text
        
        Args:
            text: Document text
            
        Returns:
            List of chapters with titles and content
        """
        chapters = []
        
        # Common chapter patterns
        chapter_patterns = [
            r'^Chapter\s+\d+',
            r'^CHAPTER\s+\d+',
            r'^\d+\.\s+[A-Z]',
            r'^Part\s+\d+',
            r'^Section\s+\d+'
        ]
        
        # Split by lines
        lines = text.split('\n')
        
        current_chapter = None
        current_content = []
        
        for i, line in enumerate(lines):
            # Check if line matches chapter pattern
            is_chapter = False
            for pattern in chapter_patterns:
                if re.match(pattern, line.strip()):
                    is_chapter = True
                    break
            
            if is_chapter:
                # Save previous chapter
                if current_chapter:
                    chapters.append({
                        'title': current_chapter,
                        'content': '\n'.join(current_content),
                        'start_line': current_start,
                        'end_line': i - 1
                    })
                
                # Start new chapter
                current_chapter = line.strip()
                current_content = []
                current_start = i
            else:
                # Add to current chapter
                if current_chapter:
                    current_content.append(line)
        
        # Save last chapter
        if current_chapter:
            chapters.append({
                'title': current_chapter,
                'content': '\n'.join(current_content),
                'start_line': current_start,
                'end_line': len(lines) - 1
            })
        
        # If no chapters found, treat whole document as one
        if not chapters:
            chapters.append({
                'title': 'Full Document',
                'content': text,
                'start_line': 0,
                'end_line': len(lines) - 1
            })
        
        logger.info(f"Extracted {len(chapters)} chapters")
        return chapters